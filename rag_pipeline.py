# -- coding: utf-8 --

import torch
import os
import json
import warnings
import re
import shutil
import time
import uuid
from tqdm import tqdm
import requests
import PyPDF2
import pandas as pd
from bs4 import BeautifulSoup
from typing import List, Dict, Any, Tuple

# LangChain & ML Imports
from langchain.prompts import PromptTemplate # ChatPromptTemplate not strictly needed for this structure
from langchain.schema.runnable import RunnablePassthrough, RunnableLambda, RunnableParallel, RunnableConfig
from langchain.schema.output_parser import StrOutputParser
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from langchain_community.retrievers import BM25Retriever
from langchain.retrievers.document_compressors import CrossEncoderReranker
from langchain_community.cross_encoders import HuggingFaceCrossEncoder

import win32com.client
try:
    import docx
except ImportError:
    print("docx package not found. Please install it using 'pip install python-docx'")

warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=UserWarning, module='torch.utils. deel_ Državna')
warnings.filterwarnings("ignore", message="Can't initialize NVML")

SEED = 42
torch.manual_seed(SEED)
if torch.cuda.is_available(): torch.cuda.manual_seed_all(SEED)
print("Imports successful.")

# Configuration
print("--- Configuring Paths ---")
DATA_ROOT = os.path.abspath("./Askalma")
WORKING_DIR = os.path.abspath("temp")
os.makedirs(WORKING_DIR, exist_ok=True)

BOT_NAME = "AskAlma"
CREATOR_INFO = "I was built by Amartya Singh, Abhishek Bansal, and Aditya Bagri."
PURPOSE = f"I am {BOT_NAME}, an AI assistant for the IIITD college website documents."

GENERIC_DOC_CHUNK_SIZE = 1500
GENERIC_DOC_CHUNK_OVERLAP = 300
BM25_K_CANDIDATES = 30
CROSS_ENCODER_MODEL_NAME = "cross-encoder/ms-marco-MiniLM-L-6-v2"
RERANK_TOP_N = 10
HISTORY_TURNS_FOR_CONDENSING = 3
LLM_TEMPERATURE_CONDENSE = 0.0
LLM_TEMPERATURE_ANSWER = 0.0

DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
TORCH_DTYPE = torch.float16 if DEVICE == "cuda" else torch.float32
print(f"Using device for potential PyTorch ops: {DEVICE}")
print(f"Using dtype for potential PyTorch ops: {TORCH_DTYPE}")
print(f"HISTORY_TURNS_FOR_CONDENSING set to: {HISTORY_TURNS_FOR_CONDENSING}")
print(f"LLM_TEMPERATURE_CONDENSE set to: {LLM_TEMPERATURE_CONDENSE}")
print(f"LLM_TEMPERATURE_ANSWER set to: {LLM_TEMPERATURE_ANSWER}")
print("Configuration loaded.")

# --- Data Loading & Extraction Functions ---
def clean_text(text: str) -> str:
    if not isinstance(text, str): text = str(text)
    text = re.sub(r'\s+', ' ', text).lower()
    text = re.sub(r'[^a-zA-Z0-9\s.,!?-]', '', text)
    text = re.sub(r'nan', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\.{2,}', '.', text)
    return text.strip()

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            if reader.is_encrypted:
                try: reader.decrypt('')
                except: print(f"Warning: Could not decrypt PDF {os.path.basename(file_path)}"); return ""
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text: text += page_text + "\n"
                except Exception as page_e:
                    print(f"    Error extracting text from page {page_num} of PDF {os.path.basename(file_path)}: {page_e}")
    except Exception as e: print(f"    Error reading PDF {os.path.basename(file_path)}: {e}")
    return clean_text(text)

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc_obj = docx.Document(file_path)
        full_text = [para.text for para in doc_obj.paragraphs]
        text = '\n'.join(full_text)
    except Exception as e: print(f"    Error reading DOCX {os.path.basename(file_path)}: {e}")
    return clean_text(text)

def doc_to_text(doc_filepath: str, output_format="txt") -> str:
    text_content = ""; word_instance = None; com_initialized = False
    try:
        win32com.client.pythoncom.CoInitialize(); com_initialized = True
        word_instance = win32com.client.Dispatch("Word.Application"); word_instance.Visible = False
        abs_doc_filepath = os.path.abspath(doc_filepath)
        temp_dir = os.path.join(WORKING_DIR, "doc_conversion_temp"); os.makedirs(temp_dir, exist_ok=True)
        temp_output_filename = os.path.splitext(os.path.basename(doc_filepath))[0] + f"_{uuid.uuid4().hex[:6]}.txt"
        temp_output_filepath = os.path.join(temp_dir, temp_output_filename); abs_temp_output_filepath = os.path.abspath(temp_output_filepath)
        doc = None
        try:
            doc = word_instance.Documents.Open(abs_doc_filepath)
            if output_format.lower() == "txt": doc.SaveAs(abs_temp_output_filepath, FileFormat=2)
            else: print(f"Unsupported .doc format: {output_format}"); return ""
        finally:
            if doc: doc.Close(False)
        if word_instance: word_instance.Quit(); word_instance = None 
        time.sleep(0.2) 
        with open(abs_temp_output_filepath, 'r', encoding='utf-8', errors='replace') as f: text_content = f.read()
        try: os.remove(abs_temp_output_filepath)
        except OSError as e: print(f"Warning: Could not remove temp {abs_temp_output_filepath}: {e}")
        return clean_text(text_content)
    except Exception as e: print(f"    Error converting .doc '{os.path.basename(doc_filepath)}': {e}"); return ""
    finally:
        if word_instance:
            try: word_instance.Quit()
            except: pass
        if com_initialized: win32com.client.pythoncom.CoUninitialize()

def excel_to_text(xlsx_filepath: str) -> str:
    try:
        xls = pd.ExcelFile(xlsx_filepath); text_parts = []
        if not xls.sheet_names: print(f"    Warning: No sheets in {os.path.basename(xlsx_filepath)}."); return ""
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
            sheet_text_content = [" | ".join(str(cell) for cell in row if pd.notna(cell) and str(cell).strip()) 
                                  for _, row in df.iterrows() if any(pd.notna(cell) and str(cell).strip() for cell in row)]
            if sheet_text_content: text_parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_text_content))
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text) if full_text.strip() else ""
    except Exception as e: print(f"    Error converting Excel '{os.path.basename(xlsx_filepath)}': {e}"); return ""

def extract_text_from_html(file_path: str) -> str:
    text = "";
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            soup = BeautifulSoup(file, 'html.parser'); text = soup.get_text(separator='\n', strip=True)
    except Exception as e: print(f"    Error reading HTML {os.path.basename(file_path)}: {e}")
    return clean_text(text)
def extract_text_from_table(file_path: str) -> str: return extract_text_from_html(file_path)

def extract_text_from_json_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
    filename = os.path.basename(file_path)
    derived_course_code = os.path.splitext(filename)[0].upper() 
    text_parts = [f"Information from file: {filename}."] # Changed intro
    metadata_out = {"derived_course_code": derived_course_code}
    try:
        with open(file_path, 'r', encoding='utf-8') as f: data = json.load(f)
        if isinstance(data, dict) and ("Course Code" in data or "Course Name" in data):
            course_code_from_json = data.get("Course Code", derived_course_code)
            course_name = data.get("Course Name", "Unknown Course")
            metadata_out["course_code_internal"] = course_code_from_json 
            metadata_out["course_name_internal"] = course_name
            text_parts.append(f"This document pertains to course {course_code_from_json}, titled '{course_name}'. The course code is {course_code_from_json} and the name is {course_name}.")
            if data.get("Course Description"): text_parts.append(f"Description: {data['Course Description']}")
            if data.get("Credits"): text_parts.append(f"Credits: {data['Credits']}.")
            prereqs = data.get("Prerequisites", {});
            if isinstance(prereqs, dict):
                if prereqs.get("Mandatory"): text_parts.append(f"Mandatory prerequisites: {prereqs['Mandatory']}.")
                if prereqs.get("Desirable"): text_parts.append(f"Desirable prerequisites: {prereqs['Desirable']}.")
            outcomes = data.get("Course Outcomes", {});
            if isinstance(outcomes, dict) and outcomes:
                 outcome_list = [f"{v}" for k,v in outcomes.items()]; text_parts.append(f"Learning outcomes: {'; '.join(outcome_list)}.")
            remaining_details = []
            handled_keys = {"Course Code", "Course Name", "Course Description", "Credits", "Prerequisites", "Course Outcomes", "Weekly Lecture Plan"}
            for key, value in data.items():
                if key not in handled_keys and isinstance(value, (str, int, float, bool)): remaining_details.append(f"{key}: {value}")
            if remaining_details: text_parts.append(f"Additional details: {'; '.join(remaining_details)}.")
            text_parts.append(f"End of information for {course_code_from_json} ({course_name}).")
        # ... (rest of generic dict/list handling from your previous version) ...
        elif isinstance(data, dict): 
            text_parts.append("The file contains the following key-value data:")
            for key, value in data.items():
                if isinstance(value, (str,int,float,bool)): text_parts.append(f"- {key}: {value}")
        elif isinstance(data, list):
            text_parts.append("The file contains a list, including items like:")
            for i, item in enumerate(data[:3]):
                if isinstance(item, dict):
                    item_desc = "; ".join(f"{k} is {v}" for k,v in item.items() if v is not None and isinstance(v,(str,int,float,bool)))
                    text_parts.append(f"Item {i+1}: {item_desc[:100]}...")
                else: text_parts.append(f"Item {i+1}: {str(item)[:100]}...")
        else: text_parts.append(f"Raw content: {json.dumps(data)}")
        return clean_text("\n".join(filter(None, text_parts))), metadata_out
    except Exception as e: print(f"    Error processing JSON file {filename}: {e}"); return "", metadata_out

def extract_text_from_llm_description_file(file_path: str) -> Tuple[str, Dict[str, Any]]:
    filename = os.path.basename(file_path)
    derived_course_code = "UNKNOWN_COURSE"; metadata_out = {}
    course_code_match = re.match(r"([A-Z0-9]+)(_desc)?\.txt", filename, re.IGNORECASE)
    if course_code_match:
        derived_course_code = course_code_match.group(1).upper()
        metadata_out["derived_course_code"] = derived_course_code
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f: content = f.read()
        formatted_text = f"This is an LLM-generated description for course {derived_course_code} (from file {filename}).\nDescription: {content}\nEnd of description for {derived_course_code}."
        return clean_text(formatted_text), metadata_out
    except Exception as e: print(f"    Error reading LLM description file {filename}: {e}"); return "", metadata_out
print("Data loading functions defined.")

def load_and_prepare_corpus_documents(data_root: str, generic_doc_chunk_size: int, generic_doc_chunk_overlap: int) -> List[Document]:
    print(f"Loading and preparing corpus documents from data path: {data_root} with Hybrid Chunking...")
    docs_for_generic_chunking_input: List[Tuple[str, Dict]] = []
    preserved_whole_documents: List[Document] = []
    min_doc_length_words = 10
    critical_source_configs = {
        "course_json": {"dir": "course_json", "ext": ".json", "handler": extract_text_from_json_file},
        "course_explain": {"dir": "course_explain", "ext": ".txt", "handler": extract_text_from_llm_description_file}}
    generic_source_configs = {
        "attachments": {"dir": "attachments", "handlers": {'.pdf': extract_text_from_pdf, '.docx': extract_text_from_docx, '.doc': doc_to_text, '.xlsx': excel_to_text, '.xls': excel_to_text}},
        "html": {"dir": "html", "handlers": {'.html': extract_text_from_html}},
        "tables": {"dir": "tables", "handlers": {'.html': extract_text_from_table}},
        "text_pdfs": {"dir": "text_pdfs", "handlers": {'.txt': lambda fp: clean_text(open(fp, 'r', encoding='utf-8', errors='replace').read())}},
    }
    for source_key, config in critical_source_configs.items():
        dir_path = os.path.join(data_root, config["dir"])
        if not os.path.isdir(dir_path): print(f"Warning: Critical source directory not found - {dir_path}, skipping."); continue
        print(f"Processing CRITICAL source directory: {dir_path}...")
        for filename in tqdm(os.listdir(dir_path), desc=f"Loading from {source_key}"):
            if filename.lower().endswith(config["ext"]):
                file_path = os.path.join(dir_path, filename)
                text, derived_meta = config["handler"](file_path)
                if text and len(text.split()) >= min_doc_length_words:
                    base_metadata = {"source": source_key, "file": filename, "full_path": file_path}
                    combined_metadata = {**base_metadata, **derived_meta}
                    preserved_whole_documents.append(Document(page_content=text, metadata=combined_metadata))
                elif text: print(f"    Skipping short critical document: {filename}")
    for source_key, config in generic_source_configs.items():
        dir_path = os.path.join(data_root, config["dir"])
        if not os.path.isdir(dir_path): print(f"Warning: Generic source directory not found - {dir_path}, skipping."); continue
        print(f"Processing GENERIC source directory: {dir_path}...")
        for filename in tqdm(os.listdir(dir_path), desc=f"Loading from {source_key}"):
            file_path = os.path.join(dir_path, filename)
            file_actual_ext = os.path.splitext(filename)[1].lower()
            if file_actual_ext in config["handlers"]:
                text = config["handlers"][file_actual_ext](file_path)
                if text and len(text.split()) >= min_doc_length_words:
                    metadata = {"source": source_key, "file": filename, "full_path": file_path}
                    docs_for_generic_chunking_input.append((text, metadata))
                elif text: print(f"    Skipping short generic document: {filename}")
    single_factual_file = os.path.join(data_root, 'factual_data_spanbert.json')
    if os.path.exists(single_factual_file):
        print(f"Processing single factual data file for chunking: {single_factual_file}...")
        try:
            with open(single_factual_file, 'r', encoding='utf-8') as f: factual_data_list = json.load(f)
            for i, item in enumerate(tqdm(factual_data_list, desc="Extracting facts for chunking")):
                if "facts" in item and isinstance(item["facts"], str):
                    facts_text = clean_text(item["facts"])
                    if facts_text and len(facts_text.split()) >= min_doc_length_words:
                        metadata = {"source": "facts_spanbert_json", "file": os.path.basename(single_factual_file), "item_index": i, "original_item_title": item.get("title", "N/A")}
                        docs_for_generic_chunking_input.append((facts_text, metadata))
        except Exception as e: print(f"Error loading or processing single factual data file {single_factual_file}: {e}")
    else: print(f"Warning: Single factual data file not found at {single_factual_file}, skipping.")
    generic_chunks = []
    if docs_for_generic_chunking_input:
        print(f"Chunking {len(docs_for_generic_chunking_input)} generic document contents...")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=generic_doc_chunk_size, chunk_overlap=generic_doc_chunk_overlap,
            length_function=len, add_start_index=True)
        temp_generic_docs = [Document(page_content=text, metadata=meta) for text, meta in docs_for_generic_chunking_input]
        generic_chunks = text_splitter.split_documents(temp_generic_docs)
        print(f"Created {len(generic_chunks)} chunks from generic documents.")
    final_corpus_documents = preserved_whole_documents + generic_chunks
    print(f"\nPrepared a total of {len(final_corpus_documents)} documents/chunks for the corpus.")
    print(f"  Preserved whole critical documents: {len(preserved_whole_documents)}")
    print(f"  Chunks from generic documents: {len(generic_chunks)}")
    from collections import Counter
    print("Final corpus document counts by source type:", Counter(doc.metadata.get('source', 'unknown') for doc in final_corpus_documents))
    return final_corpus_documents

# --- Hybrid Retriever Initialization ---
def initialize_hybrid_retriever_components(
    all_corpus_docs: List[Document],
    bm25_k_candidates: int, cross_encoder_model_name: str, rerank_top_n: int
) -> Tuple[BM25Retriever | None, CrossEncoderReranker | None]:
    if not all_corpus_docs: raise ValueError("No corpus documents provided for hybrid retriever.")
    print(f"Initializing Hybrid Retriever (BM25 -> CrossEncoder) on {len(all_corpus_docs)} corpus items...")
    print(f"Initializing BM25Retriever with k={bm25_k_candidates}...")
    try:
        bm25_retriever = BM25Retriever.from_documents(documents=all_corpus_docs, k=bm25_k_candidates)
        print("BM25Retriever initialized.")
    except Exception as e: print(f"Error initializing BM25Retriever: {e}."); raise
    print(f"Initializing CrossEncoderReranker with model '{cross_encoder_model_name}' and top_n={rerank_top_n}...")
    try:
        cross_encoder_model_obj = HuggingFaceCrossEncoder(model_name=cross_encoder_model_name)
        reranker = CrossEncoderReranker(model=cross_encoder_model_obj, top_n=rerank_top_n)
        print("CrossEncoderReranker initialized.")
    except Exception as e: print(f"Error initializing CrossEncoderReranker: {e}."); raise
    return bm25_retriever, reranker

# --- LLM, Prompt, Chain functions ---
def load_persona_data(persona_file_path: str) -> tuple[str, dict, dict]:
    try:
        with open(persona_file_path, 'r', encoding='utf-8') as f: data = json.load(f)
        return data.get("persona", ""), data.get("faqs", {}), data.get("paths", {})
    except Exception as e: print(f"Error loading persona: {e}"); return "AI Assistant.", {}, {}
PERSONA_FILE_PATH = os.path.join(DATA_ROOT, "persona.json")

def llm_pipeline_function(chain_input_dict: Dict[str, Any], temperature_to_use: float) -> str:
    raw_full_prompt = chain_input_dict.get("full_prompt_str", "")
    chat_history: List[Dict[str, str]] = chain_input_dict.get("chat_history", [])
    full_prompt_as_string = ""
    if hasattr(raw_full_prompt, 'to_string'): full_prompt_as_string = raw_full_prompt.to_string()
    elif isinstance(raw_full_prompt, str): full_prompt_as_string = raw_full_prompt
    else: # Fallback for LCEL structure where prompt might be the dict itself
        if isinstance(chain_input_dict, dict): # Check if it's a dict directly
            # Try to find a string value that looks like a prompt
            for val in chain_input_dict.values():
                if isinstance(val, str) and "User Question:" in val and "Context Documents:" in val:
                    full_prompt_as_string = val; break
                elif hasattr(val, 'to_string'): # If a value in the dict is a PromptValue
                    val_str = val.to_string()
                    if "User Question:" in val_str and "Context Documents:" in val_str:
                        full_prompt_as_string = val_str; break
        if not full_prompt_as_string:
            print(f"Warning: 'full_prompt_str' not resolved in llm_pipeline_function from: {chain_input_dict}")
            return "Error: Internal prompt generation problem."

    _p,_,_ = load_persona_data(PERSONA_FILE_PATH)
    msgs = [{"role":"system", "content":f"{_p if _p else PURPOSE} Respond in English. Think steps inside <think></think> tags if needed, also in English."}]
    for turn in chat_history: # Pass full history to LLM for answer generation
        if "query" in turn and "response" in turn:
            msgs.append({"role": "user", "content": str(turn["query"])})
            msgs.append({"role": "assistant", "content": str(turn["response"])}) 
    msgs.append({"role": "user", "content": full_prompt_as_string}) # The RAG prompt
    
    payload = {"model":"deepseek-r1-distill-qwen-7b","messages":msgs,"temperature":temperature_to_use,"max_tokens":2048,"stream":False}
    try:
        r = requests.post("http://localhost:1234/v1/chat/completions",json=payload,timeout=180) 
        r.raise_for_status(); return r.json()["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout: print("Error: LLM API timed out."); return "AI model connection timed out."
    except requests.exceptions.RequestException as e: print(f"LLM API Error: {e}"); return f"AI model connection error: {e}"
    except Exception as e: resp_text = r.text if 'r' in locals() else 'N/A'; print(f"LLM Response Error: {e}. Response: {resp_text[:500]}..."); return "AI model response error."

CONDENSE_QUESTION_PROMPT_TEMPLATE = """Given the following conversation and a follow up question, rephrase the follow up question to be a standalone question, in its original language.
If the follow up question is already a standalone question (e.g., it does not use pronouns like "it", "that", "those" to refer to previous context, and makes sense on its own), just return it as is.
If the follow up question is a greeting or a very general statement not directly answerable from specific context (e.g. "thank you", "ok"), you can return it as is.
Make sure to incorporate relevant entities or context from the chat history into the follow up question to make it self-contained.
Only return the standalone question, without any preamble or explanation.

Chat History:
{chat_history_str}

Follow Up Input: {question}
Standalone question:"""
CONDENSE_QUESTION_PROMPT = PromptTemplate.from_template(CONDENSE_QUESTION_PROMPT_TEMPLATE)

def format_chat_history_for_condenser(chat_history: List[Dict[str, str]], turns_to_include: int) -> str:
    if not chat_history or turns_to_include <= 0: return "No prior conversation."
    history_str_parts = []
    for turn_data in chat_history[-turns_to_include:]:
        user_q = str(turn_data.get('query', ''))
        raw_ai_response = str(turn_data.get('response', ''))
        cleaned_ai_response = re.sub(r"<think>.*?</think>\s*|^\*\*Answer:\*\*\s*", "", raw_ai_response, flags=re.DOTALL | re.IGNORECASE).strip()
        cleaned_ai_response = re.sub(r'<[^>]+>', '', cleaned_ai_response).strip()
        if user_q: history_str_parts.append(f"Human: {user_q}")
        if cleaned_ai_response: history_str_parts.append(f"Assistant: {cleaned_ai_response}")
    return "\n".join(history_str_parts) if history_str_parts else "No relevant prior conversation."

def clean_llm_condensed_query(llm_raw_output: str) -> str:
    # Remove <think> blocks and get the last meaningful line
    cleaned = re.sub(r"<think>.*?</think>\s*", "", llm_raw_output, flags=re.DOTALL).strip()
    # The LLM might output "Standalone question: What will I learn in the Discrete Mathematics course?"
    # We want to extract just the part after "Standalone question:"
    # Or if it just returns the question, that's fine too.
    prefix_match = re.match(r"standalone question:\s*(.*)", cleaned, re.IGNORECASE | re.DOTALL)
    if prefix_match:
        return prefix_match.group(1).strip()
    
    # Fallback: if no "Standalone question:" prefix, return the cleaned output as is.
    # This handles cases where the LLM directly returns the rephrased question.
    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    return lines[-1] if lines else cleaned # Return last non-empty line or the cleaned string

def create_rag_prompt(bot_name: str, creator_info: str, purpose: str) -> PromptTemplate:
    template_str = f"""You are {bot_name}, a helpful AI assistant for the IIITD college website.
Your persona is friendly, knowledgeable about IIITD based only on the provided context, and strictly focused on assisting with IIITD-related queries.
{creator_info}
{purpose}

**Instructions & Guardrails:**
1.  **Respond in English Only:** All parts of your response, including any internal thought processes or reasoning steps (like those within <think></think> tags if you use them), MUST be in English.
2.  **Prioritize Context:** Base your answers *exclusively* on the provided "Context Documents" below.
3.  **Acknowledge Limits:** If the context does not contain the answer, or is empty/irrelevant, clearly state "Based on the available IIITD documents, I don't have specific information about that." Do not try to answer from general knowledge if context is missing or irrelevant.
4.  **Refuse Out-of-Scope:** Politely decline for requests outside IIITD website documents.
5.  **Goal:** Provide accurate, concise answers based *only* on the context. If the context doesn't directly answer, say so.
6.  **Understanding the Query:** The user's original question was: "{{original_question}}". For clarity in document retrieval, this was interpreted or rephrased as: "{{rewritten_query_for_context}}". Your answer should address the user's original intent as reflected in the rephrased question, using the provided context.

**Context Documents:**
{{context}}

**Rephrased Question (used for finding context):**
{{rewritten_query_for_context}}

**(Answer the Rephrased Question based on the Context Documents, keeping in mind the user's original phrasing if helpful for nuance.)**
Answer ({bot_name}):
""" 
    return PromptTemplate(
        input_variables=["context", "original_question", "rewritten_query_for_context", "chat_history"], 
        template=template_str
    )

def format_docs(docs: List[Document]) -> str:
    if not docs: return "No relevant documents found."
    return "\n\n".join([doc.page_content for doc in docs if isinstance(doc, Document)])


def get_rag_chain(
    bm25_retriever: BM25Retriever, 
    reranker: CrossEncoderReranker,
    main_rag_prompt_template: PromptTemplate, # Renamed for clarity
    condense_q_prompt: PromptTemplate,    # Renamed for clarity
    hist_turns_for_condensing: int      # Renamed for clarity
):
    # --- Sub-chain for Query Condensing ---
    def clean_condensed_query(llm_output: str) -> str:
        cleaned = re.sub(r"<think>.*?</think>\s*", "", llm_output, flags=re.DOTALL).strip()
        prefix_match = re.match(r"standalone question:\s*(.*)", cleaned, re.IGNORECASE | re.DOTALL)
        if prefix_match: return prefix_match.group(1).strip()
        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
        return lines[-1] if lines else cleaned

    prepare_condense_inputs = RunnableParallel({
        "question": RunnableLambda(lambda x: x["question"], name="OriginalQuestionForCondenser"),
        "chat_history_str": RunnableLambda(lambda x: format_chat_history_for_condenser(x.get("chat_history", []), hist_turns_for_condensing), name="FormatHistoryForCondenser")
    })
    condense_question_sub_chain = (
        prepare_condense_inputs
        | condense_q_prompt
        | RunnableLambda(lambda x_prompt_value: {"full_prompt_str": x_prompt_value.to_string(), "chat_history": []}, name="PrepareLLMInputForCondenser")
        | RunnableLambda(lambda x_dict: llm_pipeline_function(x_dict, temperature_to_use=LLM_TEMPERATURE_CONDENSE), name="CondenserLLMCall")
        | RunnableLambda(clean_condensed_query, name="CleanCondenserOutput")
    )
    # --- End of Condense Question Sub-Chain ---

    # --- Sub-chain for Retrieval and Reranking ---
    def perform_bm25_then_rerank_func(payload: Dict[str, Any]) -> List[Document]:
        query_for_retrieval_str = payload["query_for_retrieval"] 
        print(f"\n[DEBUG RAG_CHAIN] Stage: BM25 Retrieval")
        print(f"[DEBUG RAG_CHAIN]   Query for BM25 (Rewritten): \"{query_for_retrieval_str}\"")
        bm25_candidate_docs = bm25_retriever.get_relevant_documents(query_for_retrieval_str)
        print(f"[DEBUG RAG_CHAIN]   BM25 Retrieved {len(bm25_candidate_docs)} docs (Top {BM25_K_CANDIDATES}). First 3:") # Use config
        for i, doc in enumerate(bm25_candidate_docs[:3]):
            print(f"    BM25 Cand {i+1} (File: {doc.metadata.get('file','N/A')}): {doc.page_content[:70].replace(os.linesep,' ')}...")
        if not bm25_candidate_docs: print("[DEBUG RAG_CHAIN]   BM25 returned no candidates."); return []
        print(f"[DEBUG RAG_CHAIN] Stage: CrossEncoder Reranking (Top {RERANK_TOP_N} from {len(bm25_candidate_docs)})") # Use config
        print(f"[DEBUG RAG_CHAIN]   Query for Reranker (Rewritten): \"{query_for_retrieval_str}\"")
        reranked_docs = reranker.compress_documents(documents=bm25_candidate_docs, query=query_for_retrieval_str)
        print(f"[DEBUG RAG_CHAIN]   CrossEncoder Reranked to {len(reranked_docs)} docs. First {RERANK_TOP_N}:") # Use config
        for i, doc in enumerate(reranked_docs): # Log all reranked docs
             print(f"    Reranked Doc {i+1} (File: {doc.metadata.get('file','N/A')}): {doc.page_content[:70].replace(os.linesep,' ')}...")
        return reranked_docs

    retrieve_and_rerank_docs_chain = (
        RunnablePassthrough.assign( 
            query_for_retrieval=condense_question_sub_chain 
        ) 
        | RunnableLambda(perform_bm25_then_rerank_func, name="RetrieveAndRerankDocuments")
    )
    # --- End of Retrieval and Reranking Sub-Chain ---
    
    # Prepare final inputs for the main RAG prompt
    # This takes the original chain input {"question": ..., "chat_history": ...}
    final_prompt_inputs_generator = RunnableParallel({
        "context": (retrieve_and_rerank_docs_chain | RunnableLambda(format_docs, name="FormatRetrievedDocs")),
        "original_question": RunnableLambda(lambda x: x["question"]),                 # CORRECTED
        "rewritten_query_for_context": condense_question_sub_chain, # Re-run to get rewritten query for the prompt
        "chat_history": RunnableLambda(lambda x: x.get("chat_history", []))       # CORRECTED
    })
    
    prepare_for_final_llm_call = {
        "full_prompt_str": main_rag_prompt_template,
        "chat_history": RunnableLambda(lambda x: x["chat_history"]) # CORRECTED (gets from output of final_prompt_inputs_generator)
    }

    chain = (
        final_prompt_inputs_generator 
        | prepare_for_final_llm_call
        | RunnableLambda(lambda x_dict: llm_pipeline_function(x_dict, temperature_to_use=LLM_TEMPERATURE_ANSWER), name="AnswerGenerationLLMCall")
        | StrOutputParser()
    )
    return chain


def initialize_knowledge_base(
    data_root_path: str,
    generic_doc_cs_param: int, generic_doc_co_param: int,
    bm25_k_candidates_param: int, cross_encoder_model_name_param: str, rerank_top_n_param: int,
    persona_file_path_kb: str
) -> Tuple[BM25Retriever | None, CrossEncoderReranker | None, str, dict, dict, PromptTemplate | None, PromptTemplate | None]:
    print("\n--- Initializing Knowledge Base (Hybrid Chunking & LLM Condense -> BM25 -> CrossEncoder) ---")
    corpus_docs = load_and_prepare_corpus_documents(
        data_root_path,
        generic_doc_chunk_size=generic_doc_cs_param,
        generic_doc_chunk_overlap=generic_doc_co_param
    )
    if not corpus_docs: print("No corpus documents prepared."); return None, None, "", {}, {}, None, None
    try:
        bm25_retriever_obj, reranker_obj = initialize_hybrid_retriever_components(
            all_corpus_docs=corpus_docs,
            bm25_k_candidates=bm25_k_candidates_param,
            cross_encoder_model_name=cross_encoder_model_name_param,
            rerank_top_n=rerank_top_n_param)
    except Exception as e: print(f"Error hybrid retriever init: {e}"); import traceback; traceback.print_exc(); return None, None, "", {}, {}, None, None
    if bm25_retriever_obj is None or reranker_obj is None:
        print("Retriever/Reranker init failed."); return None, None, "", {}, {}, None, None
    persona_text, faqs_data, paths_data = load_persona_data(persona_file_path_kb)
    rag_prompt_template_obj = create_rag_prompt(BOT_NAME, CREATOR_INFO, PURPOSE)
    condense_prompt_obj = CONDENSE_QUESTION_PROMPT 
    print("--- Knowledge Base Initialized ---")
    return bm25_retriever_obj, reranker_obj, persona_text, faqs_data, paths_data, rag_prompt_template_obj, condense_prompt_obj

if __name__ == "__main__":
    bm25_retriever, reranker, persona, faqs, paths, main_rag_prompt, condense_prompt = initialize_knowledge_base(
        DATA_ROOT,
        GENERIC_DOC_CHUNK_SIZE, GENERIC_DOC_CHUNK_OVERLAP,
        BM25_K_CANDIDATES, CROSS_ENCODER_MODEL_NAME, RERANK_TOP_N,
        PERSONA_FILE_PATH
    )

    if bm25_retriever is None or reranker is None or main_rag_prompt is None or condense_prompt is None:
        print("Failed to initialize knowledge base. Exiting.")
    else:
        standard_rag_chain = get_rag_chain(
            bm25_retriever, reranker, 
            main_rag_prompt, 
            condense_prompt,
            HISTORY_TURNS_FOR_CONDENSING
        )
        print("\n--- RAG Chain with Query Condensing, Hybrid Retrieval & History Test ---")
        session_history_for_test: List[Dict[str, str]] = []
        test_conversation = [
            {"q": "What is the course code for Discrete Mathematics?", "use_hist_for_condense": 0},
            {"q": "What will I learn in it?", "use_hist_for_condense": HISTORY_TURNS_FOR_CONDENSING},
            {"q": "what is the refernce book for it ?", "use_hist_for_condense": HISTORY_TURNS_FOR_CONDENSING} 
        ]
        for i, turn_info in enumerate(test_conversation):
            query = turn_info["q"]
            hist_turns_for_this_condense = turn_info["use_hist_for_condense"]
            print(f"\n--- Test Turn {i + 1} ---")
            print(f"❓ User Query: {query}")
            condense_prompt_input_display = {
                "question": query,
                "chat_history_str": format_chat_history_for_condenser(session_history_for_test, hist_turns_for_this_condense)
            }
            print(f"🗣️  Input to Condense Prompt (debug): {json.dumps(condense_prompt_input_display, indent=2)}")
            print("-" * 20)
            try:
                print(f"💬 Generating response from {BOT_NAME}...")
                chain_input = {"question": query, "chat_history": list(session_history_for_test)}
                final_answer = standard_rag_chain.invoke(chain_input)
                print(f"\n🤖 Answer ({BOT_NAME}):")
                final_answer_cleaned = re.sub(r"<think>.*?</think>\s*|^\*\*Answer:\*\*\s*", "", final_answer, flags=re.DOTALL|re.IGNORECASE).strip()
                print(final_answer_cleaned)
                session_history_for_test.append({"query": query, "response": final_answer})
            except Exception as e:
                print(f"\n💥 ERROR processing query '{query}': {e}"); import traceback; traceback.print_exc()
            print("=" * 50)
        print("\n--- All tests completed ---")
        print("Exiting RAG pipeline test.")
